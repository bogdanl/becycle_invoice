"""
Becycle Payroll Processor

This Streamlit application processes a Codexfit classes report PDF to extract payroll records,
generates an HTML table with the extracted data, converts it to a PDF (using WeasyPrint),
and optionally merges it with an existing invoice PDF (using PyMuPDF).

Requirements:
    - streamlit
    - pandas
    - PyMuPDF
    - weasyprint

Usage:
    Run this file with Streamlit:
        streamlit run becycle_invoice.py
"""

import io
import re
from datetime import datetime

import pandas as pd
import fitz  # PyMuPDF
import streamlit as st
from weasyprint import HTML
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT

# Set the Streamlit page configuration.
st.set_page_config(page_title="Becycle Payroll Processor", layout="centered")

st.markdown(
    """
    <style>
    .reportview-container .main .block-container{
        max-width: 800px;
        padding: 2rem 1rem;
    }
    </style>
    """,
    unsafe_allow_html=True
)


def parse_payroll_pdf(pdf_bytes):
    """
    Parse the uploaded PDF bytes and extract payroll records.

    Each record contains:
      - date (without weekday)
      - time
      - event (renamed from Class)
      - studio (full info)
      - B, A, N (bookings, attended, no-shows)
      - total (the total payment amount)

    Args:
        pdf_bytes (bytes): PDF file bytes.

    Returns:
        list: A list of dictionaries containing payroll record data.
    """
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    day_names = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]
    records = []
    current_date = None
    current_record = None

    # Extract text lines from the PDF while preserving reading order.
    lines = []
    for page in doc:
        lines.extend(page.get_text().splitlines())

    # Lines to skip (headers and labels)
    skip_lines = {
        "Time Event Type", "Time", "Event Type", "Studio",
        "B / A /", "B / A / N*", "N*",
        "Base Bonus Total", "Base", "Bonus", "Total",
        "Bookings Attended No-shows Base Bonus Total",
        "Bookings", "Attended", "No-shows"
    }

    i = 0
    while i < len(lines):
        line = lines[i].strip()
        i += 1

        if not line or line in skip_lines:
            continue

        # Look for a date header like "Tuesday 04.02.2025".
        parts = line.split(maxsplit=1)
        if parts and parts[0] in day_names:
            date_str = parts[1] if len(parts) > 1 else ""
            if not date_str or not re.match(r"\d{1,2}\.\d{2}\.\d{4}$", date_str):
                if i < len(lines) and re.match(r"\d{1,2}\.\d{2}\.\d{4}$", lines[i].strip()):
                    date_str = lines[i].strip()
                    i += 1
            current_date = date_str
            current_record = None
            continue

        if line.startswith("Totals:"):
            # Skip totals sections.
            current_date = None
            current_record = None
            continue

        # If we are inside a date section...
        if current_date:
            # New class record when a line starts with a time stamp.
            if re.match(r"^\d{1,2}:\d{2}", line):
                if current_record and current_record.get("B") is None:
                    records.append(current_record)
                parts_time = line.split(" ", 1)
                time_str = parts_time[0]
                event_str = parts_time[1].strip() if len(parts_time) > 1 else ""
                current_record = {
                    "date": current_date,
                    "time": time_str,
                    "event": event_str,
                    "studio": "",
                    "B": None,
                    "A": None,
                    "N": None,
                    "total": None
                }
                continue

            if current_record:
                # Append additional lines to event name if they start with "(".
                if line.startswith("("):
                    current_record["event"] += " " + line
                    continue
                # If event is still empty, use the line.
                if (current_record["event"] == "" and current_record["studio"] == "" and
                        not re.search(r"\(\d+\)\s*$", line) and
                        not re.search(r"\d+\s*/\s*\d+\s*/\s*\d+", line)):
                    current_record["event"] = line
                    continue
                # If event is empty and line ends with capacity info (e.g. "(12)"),
                # split it into event and studio parts.
                if current_record["event"] == "" and re.search(r"\(\d+\)\s*$", line):
                    cap_idx = line.rfind('(')
                    if cap_idx != -1 and re.match(r"\(\d+\)\s*$", line[cap_idx:]):
                        event_part = line[:cap_idx].strip()
                        studio_part = line[cap_idx:].strip()
                        words = event_part.split()
                        duplicate = None
                        if len(words) >= 2:
                            if words[-1] == words[-2]:
                                duplicate = words[-1]
                                words = words[:-1]
                            elif words[0] == words[-1]:
                                duplicate = words[-1]
                                words = words[:-1]
                        event_part = " ".join(words)
                        if duplicate:
                            studio_part = (duplicate + " " + studio_part).strip()
                        current_record["event"] = event_part
                        current_record["studio"] = studio_part
                        continue

                # Look for the B/A/N (bookings/attended/no-shows) and payment info.
                bn_match = re.search(r"(\d+\s*/\s*\d+\s*/\s*\d+)", line)
                if bn_match:
                    idx = bn_match.start()
                    studio_text = line[:idx].strip()
                    ba_text = line[idx:].strip()
                    if studio_text:
                        current_record["studio"] = (current_record["studio"] + " " + studio_text).strip()
                    m = re.match(r"(\d+)\s*/\s*(\d+)\s*/\s*(\d+)\s*(.*)", ba_text)
                    if m:
                        current_record["B"] = int(m.group(1))
                        current_record["A"] = int(m.group(2))
                        current_record["N"] = int(m.group(3))
                        rest = m.group(4).strip()
                        # Extract the total payment; expect it after the B/A/N numbers.
                        if rest == "":
                            values = []
                            for _ in range(3):
                                if i < len(lines) and lines[i].strip().endswith("€"):
                                    val = lines[i].strip()
                                    values.append(val[:-1].strip() if val.endswith("€") else val)
                                    i += 1
                            if len(values) == 3:
                                current_record["total"] = values[2]
                        else:
                            currency_vals = re.findall(r"([\d.,]+)\s*€", rest)
                            if len(currency_vals) >= 3:
                                current_record["total"] = currency_vals[2]
                            else:
                                parts_cur = [p.strip() for p in rest.split("€") if p.strip()]
                                if len(parts_cur) >= 3:
                                    current_record["total"] = parts_cur[2]
                        records.append(current_record)
                        current_record = None
                    continue

                # Otherwise, append the line as additional studio information.
                current_record["studio"] = (current_record["studio"] + " " + line).strip() if current_record["studio"] else line
                continue

    return records


def generate_html_table(records, include_vat, invoice_number="", language="German", workshop_fees=0):
    if language == "German":
        headers = {
            "Date": "Datum",
            "Time": "Zeit",
            "Event": "Kurs",
            "Studio": "Studio",
            "Attendance": "Kursteilnehmer",
            "No shows": "Abwesend",
            "Total": "Betrag",
            "All total": "Nettobetrag"
        }
        invoice_number_label = "Rechnungsnummer: "
        vat_text = "Umsatzsteuer (19%)"
        total_vat_text = "Gesamtbetrag"
        page_title = "Klassenübersicht"
        workshop_fees_label = "Workshop Gebühren"
    else:
        # English definitions...
        headers = {
            "Date": "Date",
            "Time": "Time",
            "Event": "Event",
            "Studio": "Studio",
            "Attendance": "Attendance",
            "No shows": "No shows",
            "Total": "Total",
            "All total": "Total"
        }
        invoice_number_label = "Invoice number: "
        vat_text = "VAT (19%)"
        total_vat_text = "Total + VAT"
        page_title = "Class Summary"
        workshop_fees_label = "Workshop fees"

    # Build header HTML using floated divs.
    header_html = f"""
    <div id="page-header">
        <div id="header-title" class="page-title">{page_title}</div>
        {"<div id='invoice-number' class='invoice-number'>" + invoice_number_label + invoice_number + "</div>" if invoice_number else ""}
        <div style="clear: both;"></div>
    </div>
    """

    # Build main table HTML from records.
    data = []
    total_sum = 0.0
    for rec in records:
        ba = f"{rec['A']}/{rec['B']}" if rec["A"] is not None and rec["B"] is not None else ""
        try:
            total_val = float(rec["total"].replace(",", "."))
        except Exception:
            total_val = 0.0
        total_sum += total_val
        data.append({
            headers["Date"]: rec["date"],
            headers["Time"]: rec["time"],
            headers["Event"]: rec["event"],
            headers["Studio"]: rec["studio"],
            headers["Attendance"]: ba,
            headers["No shows"]: rec["N"],
            headers["Total"]: f"{total_val:.2f} €"
        })

    # Create the main table HTML.
    main_table = "<table><thead><tr>"
    main_table += f"<th>{headers['Date']}</th>"
    main_table += f"<th>{headers['Time']}</th>"
    main_table += f"<th>{headers['Event']}</th>"
    main_table += f"<th>{headers['Studio']}</th>"
    main_table += f"<th>{headers['Attendance']}</th>"
    main_table += f"<th>{headers['No shows']}</th>"
    main_table += f"<th>{headers['Total']}</th>"
    main_table += "</tr></thead><tbody>"
    for row in data:
        main_table += "<tr>"
        main_table += f"<td>{row[headers['Date']]}</td>"
        main_table += f"<td>{row[headers['Time']]}</td>"
        main_table += f"<td class='event'>{row[headers['Event']]}</td>"
        main_table += f"<td>{row[headers['Studio']]}</td>"
        main_table += f"<td>{row[headers['Attendance']]}</td>"
        main_table += f"<td>{row[headers['No shows']]}</td>"
        main_table += f"<td>{row[headers['Total']]}</td>"
        main_table += "</tr>"
    main_table += "</tbody></table>"

    # Build summary table.
    workshop_total = float(workshop_fees) * 10.0
    net_total = total_sum + workshop_total
    summary_table = "<table class='summary-table'><tbody>"
    if workshop_fees > 0:
        summary_table += (
            f"<tr class='total'><td colspan='6'>{workshop_fees_label}</td>"
            f"<td>{workshop_total:.2f} €</td></tr>"
        )
    summary_table += (
        f"<tr class='total'><td colspan='6'>{headers['All total']}</td>"
        f"<td>{net_total:.2f} €</td></tr>"
    )
    if include_vat:
        vat_amount = net_total * 0.19
        summary_table += f"<tr class='total'><td colspan='6'>{vat_text}</td><td>{vat_amount:.2f} €</td></tr>"
        summary_table += f"<tr class='total'><td colspan='6'>{total_vat_text}</td><td>{(net_total+vat_amount):.2f} €</td></tr>"
    summary_table += "</tbody></table>"

    html = f"""
    <html>
    <head>
        <style>
        @page {{
            margin-top: 120px;
            @top-left {{
                content: element(header-title);
            }}
            @top-right {{
                content: element(invoice-number);
            }}
        }}
        body {{
            margin: 0 20px 20px 20px;
            font-family: Arial, sans-serif;
            font-size: 10px;
        }}
        #header-title {{
            position: running(header-title);
        }}
        #invoice-number {{
            position: running(invoice-number);
        }}
        #page-header {{
            width: 100%;
        }}
        .page-title {{
            float: left;
            font-size: 18px;
            font-weight: bold;
        }}
        .invoice-number {{
            float: right;
            font-size: 12px;
            font-weight: bold;
            text-align: right;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            margin-top: 20px;
            margin-bottom: 20px;
        }}
        th, td {{
            border: 1px solid #ddd;
            padding: 6px;
            text-align: left;
        }}
        th {{
            background-color: #f2f2f2;
            font-weight: bold;
        }}
        tr {{
            page-break-inside: avoid;
            break-inside: avoid;
        }}
        tr.total {{
            font-weight: bold;
            background-color: #f9f9f9;
            white-space: nowrap;
        }}
        th:last-child, td:last-child {{
            white-space: nowrap;
        }}
        td.event {{
            display: -webkit-box;
            -webkit-line-clamp: 2;
            -webkit-box-orient: vertical;
            overflow: hidden;
            text-overflow: ellipsis;
        }}
    </style>
    </head>
    <body>
        {header_html}
        {main_table}
        {summary_table}
    </body>
    </html>
    """
    return html


def generate_docx_document(records, include_vat, invoice_number="", language="German", workshop_fees=0):
    """
    Generate a DOCX file (as bytes) from the extracted payroll records.
    """
    if language == "German":
        headers = {
            "Date": "Datum",
            "Time": "Zeit",
            "Event": "Kurs",
            "Studio": "Studio",
            "Attendance": "Kursteilnehmer",
            "No shows": "Abwesend",
            "Total": "Betrag",
            "All total": "Nettobetrag"
        }
        invoice_number_label = "Rechnungsnummer: "
        vat_text = "Umsatzsteuer (19%)"
        total_vat_text = "Gesamtbetrag"
        page_title = "Klassenübersicht"
        workshop_fees_label = "Workshop Gebühren"
    else:
        headers = {
            "Date": "Date",
            "Time": "Time",
            "Event": "Event",
            "Studio": "Studio",
            "Attendance": "Attendance",
            "No shows": "No shows",
            "Total": "Total",
            "All total": "Total"
        }
        invoice_number_label = "Invoice number: "
        vat_text = "VAT (19%)"
        total_vat_text = "Total + VAT"
        page_title = "Class summary"
        workshop_fees_label = "Workshop fees"

    doc = Document()
    doc.styles['Normal'].font.size = Pt(10)

    section = doc.sections[0]
    header = section.header
    header_table = header.add_table(rows=1, cols=2, width=Inches(6))
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    left_cell = header_table.cell(0, 0)
    left_cell.text = page_title
    left_paragraph = left_cell.paragraphs[0]
    left_run = left_paragraph.runs[0]
    left_run.font.size = Pt(16) 

    right_cell = header_table.cell(0, 1)
    if invoice_number:
        right_cell.text = f"{invoice_number_label}{invoice_number}"
        right_paragraph = right_cell.paragraphs[0]
        right_run = right_paragraph.runs[0]
        right_run.font.size = Pt(16) 

    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = headers["Date"]
    hdr_cells[1].text = headers["Time"]
    hdr_cells[2].text = headers["Event"]
    hdr_cells[3].text = headers["Studio"]
    hdr_cells[4].text = headers["Attendance"]
    hdr_cells[5].text = headers["No shows"]
    hdr_cells[6].text = headers["Total"]

    total_sum = 0.0
    for rec in records:
        row_cells = table.add_row().cells
        row_cells[0].text = rec["date"]
        row_cells[1].text = rec["time"]
        row_cells[2].text = rec["event"]
        row_cells[3].text = rec["studio"]
        attendance = f"{rec['A']}/{rec['B']}" if rec["A"] is not None and rec["B"] is not None else ""
        row_cells[4].text = attendance
        row_cells[5].text = str(rec["N"])
        try:
            total_val = float(rec["total"].replace(",", "."))
        except Exception:
            total_val = 0.0
        total_sum += total_val
        row_cells[6].text = f"{total_val:.2f} €"

    workshop_total = float(workshop_fees) * 10.0
    net_total = total_sum + workshop_total

    if workshop_fees > 0:
        doc.add_paragraph(f"{workshop_fees_label}: {workshop_total:.2f} €")
    doc.add_paragraph(f"{headers['All total']}: {net_total:.2f} €")
    if include_vat:
        vat_amount = net_total * 0.19
        doc.add_paragraph(f"{vat_text}: {vat_amount:.2f} €")
        doc.add_paragraph(f"{total_vat_text}: {net_total + vat_amount:.2f} €")

    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer.getvalue()


def generate_spreadsheet_document(records, include_vat, invoice_number="", language="German", workshop_fees=0):
    """
    Generate an Excel spreadsheet (as bytes) from the extracted records,
    including a summary with totals, VAT and workshop fees.
    """
    if language == "German":
        headers = {
            "Date": "Datum",
            "Time": "Zeit",
            "Event": "Kurs",
            "Studio": "Studio",
            "Attendance": "Kursteilnehmer",
            "No shows": "Abwesend",
            "Total": "Betrag"
        }
        workshop_fees_label = "Workshop Gebühren"
        summary_label = "Nettobetrag"
        vat_label = "Umsatzsteuer (19%)"
        total_vat_label = "Gesamtbetrag"
    else:
        headers = {
            "Date": "Date",
            "Time": "Time",
            "Event": "Event",
            "Studio": "Studio",
            "Attendance": "Attendance",
            "No shows": "No shows",
            "Total": "Total"
        }
        workshop_fees_label = "Workshop fees"
        summary_label = "Total"
        vat_label = "VAT (19%)"
        total_vat_label = "Total + VAT"
    
    data = []
    total_sum = 0.0
    for rec in records:
        ba = f"{rec['A']}/{rec['B']}" if rec["A"] is not None and rec["B"] is not None else ""
        try:
            total_val = float(rec["total"].replace(",", "."))
        except Exception:
            total_val = 0.0
        total_sum += total_val
        data.append({
            headers["Date"]: rec["date"],
            headers["Time"]: rec["time"],
            headers["Event"]: rec["event"],
            headers["Studio"]: rec["studio"],
            headers["Attendance"]: ba,
            headers["No shows"]: rec["N"],
            headers["Total"]: total_val
        })
    df = pd.DataFrame(data)
    
    workshop_total = float(workshop_fees) * 10.0
    net_total = total_sum + workshop_total
    
    if include_vat:
        vat_amount = net_total * 0.19
        total_vat = net_total + vat_amount

    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        df.to_excel(writer, index=False, sheet_name="Invoice")
        workbook  = writer.book
        worksheet = writer.sheets["Invoice"]

        # Define a format for merged cells
        merge_format = workbook.add_format({
            'bold': True, 
            'border': 1, 
            'align': 'left', 
            'valign': 'vcenter', 
            'font_size': 10
        })
        # Format for numeric cells:
        num_format = workbook.add_format({
            'border': 1, 
            'num_format': '0.00', 
            'font_size': 10,
            'align': 'right',
            'valign': 'vcenter'
        })

        # Determine start row after the data (leave one blank row)
        start_row = len(df) + 2
        
        # Workshop fees row (if any)
        current_row = start_row
        if workshop_fees > 0:
            worksheet.merge_range(current_row, 0, current_row, 5, workshop_fees_label, merge_format)
            worksheet.write_number(current_row, 6, workshop_total, num_format)
            current_row += 1
        
        # Net total row
        worksheet.merge_range(current_row, 0, current_row, 5, summary_label, merge_format)
        worksheet.write_number(current_row, 6, net_total, num_format)
        current_row += 1
        
        if include_vat:
            # VAT row
            worksheet.merge_range(current_row, 0, current_row, 5, vat_label, merge_format)
            worksheet.write_number(current_row, 6, vat_amount, num_format)
            current_row += 1
            # Total with VAT row
            worksheet.merge_range(current_row, 0, current_row, 5, total_vat_label, merge_format)
            worksheet.write_number(current_row, 6, total_vat, num_format)
            current_row += 1
        
    output.seek(0)
    return output.getvalue()


def main():
    if "pdf_bytes" not in st.session_state:
        st.session_state.pdf_bytes = None
    if "doc_bytes" not in st.session_state:
        st.session_state.doc_bytes = None
    if "xls_bytes" not in st.session_state:
        st.session_state.xls_bytes = None

    st.markdown(
        """
        <style>
            div[data-testid="stFileUploader"] label,
            div[data-testid="stFileUploader"] p {
                font-size: 18px !important;
            }
            .stDownloadButton button {
                font-size: 18px;
                padding: 16px 28px;
                background-color: #0078D7;
                color: white !important;
                border: none;
                border-radius: 5px;
                cursor: pointer;
            }
            .stDownloadButton button:hover, .stDownloadButton button:active {
                color: white;
            }
        </style>
        """,
        unsafe_allow_html=True
    )
    st.title("Becycle Payroll Processor")
    st.write("The uploaded or generated data is not stored or sent anywhere. You're the only one with access to it.")
    st.write("Please check that the total amounts are correct before sending the invoice.")

    language = st.selectbox("Language", ["German", "English"], index=0)
    invoice_number = st.text_input("Invoice number (optional)")

    uploaded_classes_report = st.file_uploader(
        "Upload a Codexfit report to generate a summary of the classes in various formats.",
        type="pdf"
    )
    workshop_fees = st.number_input("Number of workshop fees", min_value=0, value=0, step=1)
    include_vat = st.checkbox("Include VAT (19%)", value=False)

    if st.button("Generate invoice"):
        if uploaded_classes_report is None:
            st.error("Please upload the Classes Report PDF.")
        else:
            with st.spinner("Generating invoice, please wait..."):
                try:
                    classes_pdf_bytes = uploaded_classes_report.read()
                    records = parse_payroll_pdf(classes_pdf_bytes)
                except Exception as e:
                    st.error(f"Error processing Classes Report PDF: {e}")
                    return

                if records:
                    html_content = generate_html_table(
                        records, include_vat, invoice_number=invoice_number,
                        language=language, workshop_fees=workshop_fees
                    )
                    table_pdf = io.BytesIO()
                    HTML(string=html_content).write_pdf(table_pdf)
                    table_pdf.seek(0)
                    try:
                        doc_table = fitz.open(stream=table_pdf.getvalue(), filetype="pdf")
                        merged_pdf = fitz.open()
                        merged_pdf.insert_pdf(doc_table)
                        st.session_state.pdf_bytes = merged_pdf.write()
                        merged_pdf.close()

                        st.session_state.doc_bytes = generate_docx_document(
                            records, include_vat, invoice_number=invoice_number,
                            language=language, workshop_fees=workshop_fees
                        )
                        st.session_state.xls_bytes = generate_spreadsheet_document(
                            records, include_vat, invoice_number=invoice_number,
                            language=language, workshop_fees=workshop_fees
                        )
                    except Exception as e:
                        st.error(f"Error merging PDFs: {e}")
                        return
                else:
                    st.error("No payroll data extracted from Classes Report. Please check the PDF format.")

    st.markdown("<br>", unsafe_allow_html=True)

    # Place download buttons at the bottom of the page.
    if st.session_state.pdf_bytes and st.session_state.doc_bytes and st.session_state.xls_bytes:
        col1, col2, col3 = st.columns(3)
        with col1:
            st.download_button(
                label="Download PDF",
                data=st.session_state.pdf_bytes,
                file_name="Becycle_invoice_" + datetime.now().strftime("%d-%m-%Y") + ".pdf",
                mime="application/pdf",
                key="pdf_download"
            )
        with col2:
            st.download_button(
                label="Download DOC",
                data=st.session_state.doc_bytes,
                file_name="Becycle_invoice_" + datetime.now().strftime("%d-%m-%Y") + ".docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="doc_download"
            )
        with col3:
            st.download_button(
                label="Download Spreadsheet",
                data=st.session_state.xls_bytes,
                file_name="Becycle_invoice_" + datetime.now().strftime("%d-%m-%Y") + ".xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                key="xls_download"
            )

if __name__ == "__main__":
    main()
